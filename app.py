import streamlit as st
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# -----------------------------
# 页面配置
# -----------------------------
st.set_page_config(page_title="🛒 电商知识测验", layout="wide")
st.title("🛒 AI云学堂")

# 🔗 大模型接口地址
BACKEND_URL = "http://localhost:8889/api/user/user/aliyun/chat"

# -----------------------------
# 试卷库
# -----------------------------
EXAM_CATEGORIES = {
    "📦 电子商务类型": {
        "试卷一：电商模式基础": [
            {"question": "以下哪种模式代表企业对消费者？", "options": ["B2B", "B2C", "C2C", "O2O"], "correct": "B2C"},
            {"question": "C2C模式的典型代表平台是？", "options": ["天猫", "京东", "淘宝", "唯品会"], "correct": "淘宝"},
            {"question": "O2O模式是指什么？", "options": ["线上到线下", "企业到企业", "消费者到消费者", "批发到零售"],
             "correct": "线上到线下"},
            {"question": "B2B平台主要服务对象是？", "options": ["个人消费者", "企业客户", "政府机构", "非营利组织"],
             "correct": "企业客户"},
            {"question": "跨境电商属于哪种业务类型？", "options": ["国内零售", "国际贸易", "本地服务", "社区团购"],
             "correct": "国际贸易"}
        ],
        "试卷二：新兴电商模式": [
            {"question": "社交电商的核心特征是什么？", "options": ["价格最低", "社交裂变", "物流最快", "商品最全"],
             "correct": "社交裂变"},
            {"question": "直播电商属于哪种营销方式？", "options": ["传统广告", "内容营销", "搜索引擎", "线下推广"],
             "correct": "内容营销"},
            {"question": "社区团购的主要优势是？", "options": ["价格贵", "配送慢", "集中采购降低成本", "商品少"],
             "correct": "集中采购降低成本"},
            {"question": "私域流量电商指的是？",
             "options": ["公共平台流量", "自有客户资源", "付费广告流量", "搜索引擎流量"], "correct": "自有客户资源"},
            {"question": "S2B2C模式中S指的是？", "options": ["供应链平台", "个人卖家", "大型商场", "物流公司"],
             "correct": "供应链平台"}
        ]
    },
    "🎯 电商运营": {
        "试卷一：运营基础知识": [
            {"question": "'双十一'购物节最初是由哪家公司发起的？", "options": ["京东", "拼多多", "阿里巴巴", "苏宁"],
             "correct": "阿里巴巴"},
            {"question": "在电商平台中，'SKU'是指什么？",
             "options": ["库存保有单位", "销售关键指标", "标准采购数量", "商品分类编码"], "correct": "库存保有单位"},
            {"question": "电商运营中的'GMV'指的是？", "options": ["毛利率", "商品交易总额", "客单价", "转化率"],
             "correct": "商品交易总额"},
            {"question": "用户运营的核心指标不包括？",
             "options": ["活跃用户数", "用户留存率", "仓库面积", "用户生命周期价值"], "correct": "仓库面积"},
            {"question": "电商平台的转化率指的是？",
             "options": ["访问到购买的比例", "点击到访问的比例", "浏览到收藏的比例", "加购到支付的比例"],
             "correct": "访问到购买的比例"}
        ],
        "试卷二：营销与推广": [
            {"question": "以下哪项是社交电商的典型特征？",
             "options": ["通过搜索引擎引流", "依赖电视广告推广", "利用社交媒体分享和裂变", "主要依靠线下门店"],
             "correct": "利用社交媒体分享和裂变"},
            {"question": "电商平台的'CPC'广告模式是指？",
             "options": ["按展示付费", "按点击付费", "按成交付费", "按时长付费"], "correct": "按点击付费"},
            {"question": "内容营销最重要的是？", "options": ["广告投放量", "价格折扣力度", "优质内容创作", "平台流量"],
             "correct": "优质内容创作"},
            {"question": "私域流量运营的主要工具是？", "options": ["百度搜索", "微信社群", "电视广告", "报纸宣传"],
             "correct": "微信社群"},
            {"question": "用户增长的AARRR模型中，第一个A代表？", "options": ["激活", "获取", "推荐", "留存"],
             "correct": "获取"}
        ]
    },
    "💻 Java基础": {
        "试卷一：Java核心概念": [
            {"question": "Java的核心特性不包括？", "options": ["面向对象", "平台无关性", "自动内存管理", "硬件直接操作"],
             "correct": "硬件直接操作"},
            {"question": "JVM的全称是？",
             "options": ["Java Virtual Machine", "Java Value Method", "Java Version Manager", "Java Variable Memory"],
             "correct": "Java Virtual Machine"},
            {"question": "Java中的'=='和'equals()'的区别是？",
             "options": ["完全相同", "==比较引用，equals比较内容", "==比较内容，equals比较引用", "都比较内容"],
             "correct": "==比较引用，equals比较内容"},
            {"question": "Java的访问修饰符中权限最大的是？", "options": ["private", "protected", "default", "public"],
             "correct": "public"},
            {"question": "Java中哪个关键字用于继承？", "options": ["implements", "extends", "inherits", "super"],
             "correct": "extends"}
        ],
        "试卷二：Java面向对象": [
            {"question": "面向对象的三大特性不包括？", "options": ["封装", "继承", "多态", "编译"], "correct": "编译"},
            {"question": "Java中的接口使用哪个关键字？", "options": ["class", "interface", "abstract", "implements"],
             "correct": "interface"},
            {"question": "Java中的抽象类可以被实例化吗？", "options": ["可以", "不可以", "有时可以", "取决于编译器"],
             "correct": "不可以"},
            {"question": "重写（Override）和重载（Overload）的区别？",
             "options": ["完全相同", "重写是子类覆盖父类方法，重载是同名不同参数",
                         "重载是子类覆盖父类方法，重写是同名不同参数", "没有区别"],
             "correct": "重写是子类覆盖父类方法，重载是同名不同参数"},
            {"question": "Java中哪个关键字可以阻止类被继承？", "options": ["static", "final", "private", "abstract"],
             "correct": "final"}
        ]
    }
}

# 创建左右布局
col1, col2 = st.columns([3, 4])

# -----------------------------
# 左侧：测验区
# -----------------------------
with col1:
    st.subheader("📋 测验区")

    # 初始化 session state
    if "selected_category" not in st.session_state:
        st.session_state.selected_category = None
    if "selected_exam" not in st.session_state:
        st.session_state.selected_exam = None
    if "current_questions" not in st.session_state:
        st.session_state.current_questions = []
    if "exam_submitted" not in st.session_state:
        st.session_state.exam_submitted = False
    if "exam_score" not in st.session_state:
        st.session_state.exam_score = 0
    if "correct_count" not in st.session_state:
        st.session_state.correct_count = 0
    if "total_count" not in st.session_state:
        st.session_state.total_count = 0

    # 选择分类
    st.markdown("### 📚 选择考试分类")
    category = st.selectbox(
        "请选择分类：",
        options=list(EXAM_CATEGORIES.keys()),
        key="category_select"
    )

    # 选择试卷
    if category:
        st.markdown("### 📝 选择试卷")
        exam_name = st.selectbox(
            "请选择试卷：",
            options=list(EXAM_CATEGORIES[category].keys()),
            key="exam_select"
        )

        # 加载试卷按钮
        if st.button("📂 加载试卷", type="primary", use_container_width=True):
            st.session_state.selected_category = category
            st.session_state.selected_exam = exam_name
            st.session_state.current_questions = EXAM_CATEGORIES[category][exam_name]
            st.session_state.exam_submitted = False  # 重置提交状态
            st.success(f"✅ 已加载：{exam_name}")
            st.rerun()

    # 显示试卷内容
    if st.session_state.current_questions:
        st.markdown("---")
        st.markdown(f"### 📄 当前试卷：{st.session_state.selected_exam}")
        st.info(f"📌 分类：{st.session_state.selected_category}")

        with st.form("quiz_form"):
            user_answers = []
            for i, q in enumerate(st.session_state.current_questions):
                st.markdown(f"**{i + 1}. {q['question']}**")
                answer = st.radio(
                    "请选择答案：",
                    options=q["options"],
                    key=f"q{i}",
                    label_visibility="collapsed"
                )
                user_answers.append({"question": q["question"], "answer": answer, "correct": q["correct"]})
                st.markdown("")  # 空行

            submitted = st.form_submit_button("✅ 提交答卷", type="primary", use_container_width=True)

        if submitted:
            # 计算分数
            correct_count = sum(1 for qa in user_answers if qa["answer"] == qa["correct"])
            total_count = len(user_answers)
            score = (correct_count / total_count) * 100

            # 保存到 session state
            st.session_state.exam_submitted = True
            st.session_state.exam_score = score
            st.session_state.correct_count = correct_count
            st.session_state.total_count = total_count

            st.markdown("---")
            st.markdown("### 📊 测验结果")
            if score >= 80:
                st.success(f"🎉 恭喜！得分：{score:.0f}分 ({correct_count}/{total_count}题正确)")
            elif score >= 60:
                st.warning(f"👍 及格！得分：{score:.0f}分 ({correct_count}/{total_count}题正确)")
            else:
                st.error(f"💪 继续努力！得分：{score:.0f}分 ({correct_count}/{total_count}题正确)")

# -----------------------------
# 右侧：AI 聊天区
# -----------------------------
with col2:
    st.subheader("🤖 智能助教")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    if "ai_summary_content" not in st.session_state:
        st.session_state.ai_summary_content = None

    if "user_answers_saved" not in st.session_state:
        st.session_state.user_answers_saved = []

    # 🔍 提交后自动生成AI总结
    if submitted and st.session_state.current_questions:
        # 清除之前的AI消息
        st.session_state.messages = [m for m in st.session_state.messages if m.get("type") != "summary"]

        with st.spinner("🧠 正在生成 AI 分析..."):
            # 保存用户答案
            st.session_state.user_answers_saved = user_answers

            # 构建提示词
            prompt_content = f"【考试分类】{st.session_state.selected_category}\n【试卷名称】{st.session_state.selected_exam}\n【得分】{score:.0f}分\n\n"
            prompt_content += "【答题详情】\n"
            for i, qa in enumerate(user_answers, 1):
                prompt_content += f"{i}. {qa['question']}\n   您的回答：{qa['answer']} {status}\n\n"
                prompt = (
                    f"你是一位资深考试教练，拥有多年学科辅导与应试策略指导经验。请基于以下学生的测验情况，提供专业、精准且富有激励性的个性化反馈。\n\n"
                    f"{prompt_content}\n\n"
                    f"请严格遵循以下要求：\n"
                    f"1. **整体评价**：结合得分水平，给予恰当的肯定或建设性鼓励，避免空泛表扬或过度批评。\n"
                    f"2. **错题诊断**：针对每道错题（或错误类型），明确指出所涉及的核心知识点，并简要分析错误原因（如概念混淆、审题偏差、计算失误、策略不当等）。\n"
                    f"3. **学习建议**：提供1–2条具体、可操作的改进建议（例如：强化某类题型训练、回归教材某章节、使用错题本复盘、提升时间分配策略等）。\n"
                    f"4. **语气风格**：专业、温暖、坚定，体现教练式引导——既有高标准，又传递信心。\n"
                    f"5. **输出格式**：分三部分呈现——【整体反馈】、【错题与知识点分析】、【后续行动建议】，语言简洁，总字数控制在200字以内。"
                )

                try:
                    response = requests.post(
                        BACKEND_URL,
                        headers={"Content-Type": "application/json"},
                        json={"prompt": prompt},
                        timeout=30
                    )
                    data = response.json()
                    ai_summary = data.get("response", "AI 未返回有效内容。")

                    # 保存AI总结内容
                    st.session_state.ai_summary_content = ai_summary

                    st.session_state.messages.append({"role": "ai", "content": ai_summary, "type": "summary"})
                except Exception as e:
                    st.session_state.messages.append({"role": "ai", "content": f"⚠️ AI 分析失败：{str(e)}", "type": "error"})

            status = "✅ 正确" if qa["answer"] == qa["correct"] else f"❌ 错误（正确答案：{qa['correct']}）"

    # ✅ 生成 DOCX 文档的函数
    def create_docx():
        doc = Document()

        # 文档标题
        title = doc.add_heading('电子商务知识测验 - AI 分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成时间
        time_para = doc.add_paragraph()
        time_run = time_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = RGBColor(128, 128, 128)
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 考试信息
        doc.add_heading('📋 考试信息', level=1)
        info_para = doc.add_paragraph()
        info_para.add_run(f'考试分类：{st.session_state.selected_category}\n').font.size = Pt(11)
        info_para.add_run(f'试卷名称：{st.session_state.selected_exam}\n').font.size = Pt(11)

        score_run = info_para.add_run(
            f'测验得分：{st.session_state.exam_score:.0f}分 ({st.session_state.correct_count}/{st.session_state.total_count}题正确)')
        score_run.font.size = Pt(12)
        score_run.bold = True
        score_run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph('=' * 60)

        # 答题详情
        doc.add_heading('📝 答题详情', level=1)

        for i, qa in enumerate(st.session_state.user_answers_saved, 1):
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f'{i}. {qa["question"]}')
            q_run.font.size = Pt(11)
            q_run.bold = True

            a_para = doc.add_paragraph()
            a_run = a_para.add_run(f'   您的回答：{qa["answer"]}')
            a_run.font.size = Pt(10)

            if qa["answer"] == qa["correct"]:
                a_run.font.color.rgb = RGBColor(0, 128, 0)
                status_run = a_para.add_run(' ✅ 正确')
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                a_run.font.color.rgb = RGBColor(255, 0, 0)
                status_run = a_para.add_run(f' ❌ 错误（正确答案：{qa["correct"]}）')
                status_run.font.color.rgb = RGBColor(255, 0, 0)

            doc.add_paragraph()

        doc.add_paragraph('=' * 60)

        # AI 反馈
        doc.add_heading('🤖 AI 个性化反馈', level=1)

        feedback_paragraphs = st.session_state.ai_summary_content.split('\n')
        for para_text in feedback_paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text)
                p.paragraph_format.line_spacing = 1.5
                for run in p.runs:
                    run.font.size = Pt(11)

        doc.add_paragraph()
        doc.add_paragraph('=' * 60)

        # 结束语
        footer = doc.add_paragraph()
        footer_run = footer.add_run('感谢您使用电商知识测验系统！')
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio


    # ✅ 下载按钮区域
    if st.session_state.ai_summary_content and st.session_state.exam_submitted:
        st.markdown("---")
        st.markdown("### 📥 下载分析报告")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # TXT 格式
        txt_content = f"""========================================
电子商务知识测验 - AI 分析报告
========================================

生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}

考试信息：
----------------------------------------
考试分类：{st.session_state.selected_category}
试卷名称：{st.session_state.selected_exam}
测验得分：{st.session_state.exam_score:.0f}分 ({st.session_state.correct_count}/{st.session_state.total_count}题正确)

答题详情：
----------------------------------------

"""
        for i, qa in enumerate(st.session_state.user_answers_saved, 1):
            status = "✅ 正确" if qa["answer"] == qa["correct"] else f"❌ 错误（正确答案：{qa['correct']}）"
            txt_content += f"{i}. {qa['question']}\n   您的回答：{qa['answer']} {status}\n\n"

        txt_content += f"""----------------------------------------
AI 个性化反馈：
----------------------------------------

{st.session_state.ai_summary_content}

========================================
感谢您使用电商知识测验系统！
========================================
"""

        # Markdown 格式
        md_content = f"""# 电子商务知识测验 - AI 分析报告

**生成时间：** {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}

## 📋 考试信息

- **考试分类：** {st.session_state.selected_category}
- **试卷名称：** {st.session_state.selected_exam}
- **测验得分：** {st.session_state.exam_score:.0f}分 ({st.session_state.correct_count}/{st.session_state.total_count}题正确)

---

## 📝 答题详情

"""
        for i, qa in enumerate(st.session_state.user_answers_saved, 1):
            status = "✅ 正确" if qa["answer"] == qa["correct"] else f"❌ 错误（正确答案：{qa['correct']}）"
            md_content += f"### {i}. {qa['question']}\n\n**您的回答：** {qa['answer']} {status}\n\n"

        md_content += f"""---

## 🤖 AI 个性化反馈

{st.session_state.ai_summary_content}

---

*感谢您使用电商知识测验系统！*
"""

        # 三个下载按钮
        col_btn1, col_btn2, col_btn3 = st.columns(3)

        with col_btn1:
            st.download_button(
                label="📄 TXT",
                data=txt_content,
                file_name=f"测验报告_{timestamp}.txt",
                mime="text/plain",
                use_container_width=True,
                type="secondary"
            )

        with col_btn2:
            st.download_button(
                label="📝 Markdown",
                data=md_content,
                file_name=f"测验报告_{timestamp}.md",
                mime="text/markdown",
                use_container_width=True,
                type="secondary"
            )

        with col_btn3:
            docx_file = create_docx()
            st.download_button(
                label="📑 DOCX",
                data=docx_file,
                file_name=f"测验报告_{timestamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )

        st.markdown("---")

    # 样式
    st.markdown("""
        <style>
        .chat-container {
            height: 380px;
            overflow-y: auto;
            overflow-x: hidden;
            background-color: #ffffff;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
            padding: 16px;
            margin-bottom: 16px;
            scroll-behavior: smooth;
        }
        .chat-container::-webkit-scrollbar {
            width: 8px;
        }
        .chat-container::-webkit-scrollbar-track {
            background: #f5f5f5;
            border-radius: 10px;
        }
        .chat-container::-webkit-scrollbar-thumb {
            background: #c0c0c0;
            border-radius: 10px;
        }
        .msg-user {
            background: linear-gradient(135deg, #DCF8C6 0%, #C8E6C9 100%);
            color: #000;
            border-radius: 18px 18px 4px 18px;
            padding: 12px 16px;
            margin: 10px 0 10px auto;
            max-width: 75%;
            width: fit-content;
            word-wrap: break-word;
            box-shadow: 0 2px 4px rgba(0,0,0,0.08);
            float: right;
            clear: both;
        }
        .msg-ai {
            background-color: #f9f9f9;
            color: #333;
            border: 1px solid #e8e8e8;
            border-radius: 18px 18px 18px 4px;
            padding: 12px 16px;
            margin: 10px auto 10px 0;
            max-width: 75%;
            width: fit-content;
            word-wrap: break-word;
            box-shadow: 0 2px 4px rgba(0,0,0,0.06);
            float: left;
            clear: both;
        }
        .chat-container::after {
            content: "";
            display: table;
            clear: both;
        }
        .scroll-btn-container {
            text-align: center;
            margin-top: 8px;
            margin-bottom: 12px;
        }
        .scroll-btn {
            background-color: #4CAF50;
            color: white;
            border: none;
            padding: 8px 20px;
            border-radius: 20px;
            cursor: pointer;
            font-size: 14px;
            margin: 0 5px;
            transition: all 0.3s ease;
            box-shadow: 0 2px 5px rgba(0,0,0,0.2);
        }
        .scroll-btn:hover {
            background-color: #45a049;
            box-shadow: 0 4px 8px rgba(0,0,0,0.3);
            transform: translateY(-2px);
        }
        </style>
    """, unsafe_allow_html=True)

    # 渲染聊天内容
    chat_html = '<div class="chat-container" id="chat-box">'

    if not st.session_state.messages:
        chat_html += '<div class="msg-ai">👋 您好！我是您的智能助教，完成测验后我会为您提供个性化分析。平时有任何问题都可以问我哦！</div>'

    for msg in st.session_state.messages:
        cls = "msg-user" if msg["role"] == "user" else "msg-ai"
        content = msg["content"].replace("\n", "<br>")
        chat_html += f'<div class="{cls}">{content}</div>'

    chat_html += "</div>"
    st.markdown(chat_html, unsafe_allow_html=True)

    # 滚动按钮
    st.markdown("""
        <div class="scroll-btn-container">
            <button class="scroll-btn" onclick="document.getElementById('chat-box').scrollTop = 0">⬆ 滚动到顶部</button>
            <button class="scroll-btn" onclick="document.getElementById('chat-box').scrollTop = document.getElementById('chat-box').scrollHeight">⬇ 滚动到底部</button>
        </div>
    """, unsafe_allow_html=True)

    # 自动滚动
    if st.session_state.messages:
        st.markdown("""
            <script>
            setTimeout(function() {
                const box = document.getElementById('chat-box');
                if (box) { box.scrollTop = box.scrollHeight; }
            }, 100);
            </script>
        """, unsafe_allow_html=True)

    # 聊天输入
    user_input = st.chat_input("请输入您的问题...")
    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input})
        try:
            resp = requests.post(
                BACKEND_URL,
                headers={"Content-Type": "application/json"},
                json={"prompt": user_input},
                timeout=30
            )
            data = resp.json()
            ai_reply = data.get("response", "AI 无响应。")
        except Exception as e:
            ai_reply = f"⚠️ 请求失败：{e}"
        st.session_state.messages.append({"role": "ai", "content": ai_reply})
        st.rerun()